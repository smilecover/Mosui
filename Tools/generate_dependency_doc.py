#!/usr/bin/env python3
"""生成 RightDelegateWin.qml 依赖关系分析 Word 文档 (含 Mermaid 图表)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, zlib, base64, json, requests, io

doc = Document()

# ==================== 样式设置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ==================== Mermaid 渲染工具函数 ====================
def mermaid_to_png(code: str, width: int = 1200) -> io.BytesIO | None:
    """通过 mermaid.ink / kroki.io API 渲染为 PNG"""
    try:
        # 方法1: mermaid.ink 直接 base64 (适用于中小图)
        b64 = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{b64}?type=png&width={width}"
        resp = requests.get(url, timeout=60)
        if resp.status_code == 200 and len(resp.content) > 500:
            return io.BytesIO(resp.content)

        # 方法2: 414 URL过长 → POST 到 kroki.io
        if resp.status_code == 414 or len(url) > 8000:
            print(f"  [提示] URL 过长({len(url)}chars), 改用 kroki.io POST...")
            resp2 = requests.post(
                "https://kroki.io/mermaid/png",
                data=code.encode('utf-8'),
                headers={'Content-Type': 'text/plain'},
                timeout=60
            )
            if resp2.status_code == 200 and len(resp2.content) > 500:
                return io.BytesIO(resp2.content)
            print(f"  [警告] kroki POST 失败 status={resp2.status_code}")

        # 方法3: mermaid.ink pako 压缩
        compressed = zlib.compress(code.encode('utf-8'), level=9)
        b64c = base64.urlsafe_b64encode(compressed).decode('ascii').rstrip('=')
        url3 = f"https://mermaid.ink/img/pako:{b64c}?type=png&width={width}"
        resp3 = requests.get(url3, timeout=30)
        if resp3.status_code == 200 and len(resp3.content) > 500:
            return io.BytesIO(resp3.content)

        print(f"  [警告] 所有渲染方式均失败 (mermaid.ink:{resp.status_code}, pako:{resp3.status_code})")
        return None
    except Exception as e:
        print(f"  [警告] 渲染异常: {e}")
        return None

def add_diagram(doc, title, code, w=5.8):
    """添加图表: 图片 + 源码"""
    doc.add_heading(title, level=2)
    img = mermaid_to_png(code, width=int(w * 200))
    if img:
        try:
            doc.add_picture(img, width=Inches(w))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[图片插入失败: {e}]')
    else:
        doc.add_paragraph('[图表渲染失败, 见下方源码]')
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.font.name = 'Consolas'
    doc.add_paragraph('')

# ==================== 封面 ====================
title = doc.add_heading('RightDelegateWin.qml 依赖关系分析文档', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('文件: Mosuiapp/Tools/PLC_APP/PLCPage/RightDelegateWin.qml')
doc.add_paragraph('功能: PLC液压站右侧控制面板 UI 组件')
doc.add_paragraph('')

# ==================== 1. 模块导入 ====================
doc.add_heading('一、模块导入依赖 (Imports)', level=1)
t = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['序号', '导入模块', '类型', '说明']):
    t.rows[0].cells[i].text = h
    for pp in t.rows[0].cells[i].paragraphs:
        pp.runs[0].font.bold = True
for i, d in enumerate([
    ['1', 'QtQuick', 'Qt 官方', 'QML 基础类型 (Item, Rectangle, Timer 等)'],
    ['2', 'QtQuick.Controls', 'Qt 官方', 'Qt Quick Controls 2 组件库'],
    ['3', 'QtQuick.Layouts', 'Qt 官方', 'RowLayout, ColumnLayout 布局管理器'],
    ['4', 'MosuiBasic', '项目自定义库', 'Mosui 基础组件库 (MosLabel, MosButton 等)'],
]):
    for j, v in enumerate(d):
        t.rows[i+1].cells[j].text = v

# ==================== 2. 融合总图 ====================
doc.add_heading('二、融合依赖关系总图', level=1)

doc.add_paragraph('以下单张 Mermaid 图表融合了: 组件层级树 + 数据绑定 + 文件依赖 + Repeater模型 + 事件交互 + 架构总览', style='Normal')
doc.add_paragraph('')

unified_mmd = """graph TB
    subgraph LEGEND["🟡 图例 LEGEND"]
        L1["── 实线: 组件包含/数据流向"]
        L2["- - 虚线: 主题/样式/间接依赖"]
        L3["🟠 橙色: 核心文件"]
        L4["🟢 绿色: 外部环境"]
        L5["🔵 蓝色: MosuiBasic组件"]
        L6["🟣 紫色: UI输出"]
        L7["🔴 红色: 报警/动作"]
    end

    %% ═══════════════════════════════════════════════
    %% Layer 1: 外部环境
    %% ═══════════════════════════════════════════════
    subgraph QT["Qt Framework"]
        QQ["QtQuick"]
        QQC["QtQuick.Controls"]
        QQL["QtQuick.Layouts"]
    end

    subgraph PARENT["父级窗口"]
        MAIN["PLC_APP.qml<br/>plcappwindow"]
    end

    subgraph PROPS["root Properties 属性"]
        P_YALI["yali_pressure<br/>real = 6.5"]
        P_WENDU["wendu_temperature<br/>real = 46.9"]
        P_HEAT["heatStatus<br/>int: 0/1/2"]
    end

    subgraph EXT["外部服务"]
        EXT_API["MosApi<br/>setWindowState"]
        EXT_THEME["MosTheme<br/>主题配置"]
        EXT_ICON["MosIcon<br/>图标枚举"]
        EXT_DATE["new Date()<br/>系统时钟"]
    end

    %% ═══════════════════════════════════════════════
    %% Layer 2: 核心文件 + MosuiBasic 源文件
    %% ═══════════════════════════════════════════════
    subgraph CORE["RightDelegateWin.qml (核心)"]
        direction TB
        C_IMPORT["import QtQuick, QtQuick.Controls<br/>QtQuick.Layouts, MosuiBasic"]
        C_ROOT["Item root<br/>width: 290px"]

        subgraph C_FILES["直接依赖的 MosuiBasic 文件"]
            F_LABEL["MosLabel.qml"]
            F_BUTTON["MosButton.qml"]
            F_ICONBTN["MosIconButton.qml"]
            F_CHECK["MosCheckBox.qml"]
            F_FRAME["MosFrame.qml"]
            F_GROUPBOX["MosGroupBox.qml"]
            F_SPACE["MosSpace.qml"]
            F_RECT["MosRectangle.qml"]
        end
    end

    subgraph CPP["MosuiBasic C++ 后端"]
        CPP_ICON["MosIcon.h"]
        CPP_THEME["MosTheme"]
        CPP_API["MosApi"]
    end

    C_IMPORT --> C_ROOT
    C_ROOT --> F_LABEL
    C_ROOT --> F_BUTTON
    C_ROOT --> F_ICONBTN
    C_ROOT --> F_CHECK
    C_ROOT --> F_FRAME
    C_ROOT --> F_GROUPBOX
    C_ROOT --> F_SPACE
    C_ROOT --> F_RECT

    %% ═══════════════════════════════════════════════
    %% Layer 3: 组件树 7大区域
    %% ═══════════════════════════════════════════════
    subgraph TREE["组件层级树 Component Tree"]
        direction TB

        subgraph Z1["Zone1 顶部按钮区 MosFrame"]
            Z1_Row["MosSpace RowLayout"]
            Z1_L["左侧 启动+退出"]
            Z1_R["右侧 column1 最小化+最大化"]
            Z1_B1["MosIconButton 启动"]
            Z1_B2["MosIconButton 退出→close"]
            Z1_B3["MosIconButton 最小化→MosApi"]
            Z1_B4["MosIconButton 最大化/还原"]
            Z1_Row --> Z1_L --> Z1_B1
            Z1_Row --> Z1_R --> Z1_B3
            Z1_L --> Z1_B2
            Z1_R --> Z1_B4
        end

        subgraph Z2["Zone2 时钟区 Item"]
            Z2_CL["MosLabel clockLabel<br/>yyyy/MM/dd hh:mm:ss"]
            Z2_TM["Timer 1000ms"]
        end

        subgraph Z3["Zone3 运行状态 MosFrame"]
            Z3_SL["MosLabel 正常运行<br/>color=#FFC738"]
        end

        subgraph Z4["Zone4 参数区 MosFrame"]
            Z4_CS["MosSpace ColumnLayout column2"]
            Z4_RP["Repeater x8"]
            Z4_DL["RowLayout delegate"]
            Z4_CB["MosCheckBox<br/>enabled:false"]
            Z4_SP["Item 弹性间距"]
            Z4_RC["MosRectangle<br/>50x5 红/灰"]
            Z4_CS --> Z4_RP --> Z4_DL
            Z4_DL --> Z4_CB
            Z4_DL --> Z4_SP
            Z4_DL --> Z4_RC
        end

        subgraph Z5["Zone5 按键控制 MosFrame"]
            Z5_CS["MosSpace ColumnLayout column3"]
            Z5_RP["Repeater x4"]
            Z5_BT["MosButton<br/>参数设置/硬件监控<br/>报告输出/节流阀校准"]
            Z5_CS --> Z5_RP --> Z5_BT
        end

        subgraph Z6["Zone6 液压站数据 MosGroupBox"]
            Z6_GS["MosSpace ColumnLayout"]
            Z6_R1["RowLayout row1"]
            Z6_R2["RowLayout 加热"]
            Z6_PL["MosLabel<br/>压力MPa"]
            Z6_TL["MosLabel<br/>温度C"]
            Z6_HL["Rectangle heatLight<br/>12x12 指示灯"]
            Z6_HT["MosLabel<br/>加热运行/停止/故障"]
            Z6_GS --> Z6_R1 --> Z6_PL
            Z6_GS --> Z6_R2 --> Z6_HL
            Z6_R1 --> Z6_TL
            Z6_R2 --> Z6_HT
        end

        subgraph Z7["Zone7 底部报警 MosFrame"]
            Z7_CS["MosSpace ColumnLayout column4"]
            Z7_RP["Repeater x2"]
            Z7_LB["MosLabel<br/>气源压力+743KPa<br/>油箱低液位报警"]
            Z7_CS --> Z7_RP --> Z7_LB
        end

        C_ROOT --> Z1
        C_ROOT --> Z2
        C_ROOT --> Z3
        C_ROOT --> Z4
        C_ROOT --> Z5
        C_ROOT --> Z6
        C_ROOT --> Z7
    end

    %% ═══════════════════════════════════════════════
    %% Layer 4: 数据绑定 & 事件
    %% ═══════════════════════════════════════════════
    subgraph BINDINGS["数据绑定与事件 Data Bindings"]
        direction LR

        subgraph EVENTS["事件处理"]
            EV_EXIT["onClicked 退出"]
            EV_MIN["onClicked 最小化"]
            EV_MAX["onClicked 最大化"]
            EV_TICK["Timer.onTriggered"]
        end

        subgraph ACTIONS["窗口动作"]
            ACT_CLOSE["close()"]
            ACT_MIN["showMinimized()"]
            ACT_MAX["showMaximized()"]
        end
    end

    %% 属性 → UI 绑定
    P_YALI -->|"toFixed(1)"| Z6_PL
    P_WENDU -->|"toFixed(1)"| Z6_TL
    P_HEAT -->|"switch 0灰1绿2红"| Z6_HL
    P_HEAT -->|"switch-case"| Z6_HT

    %% 时钟 → 显示
    EXT_DATE --> EV_TICK --> Z2_CL

    %% 按钮事件 → 窗口动作
    Z1_B2 --> EV_EXIT
    MAIN --> EV_EXIT --> ACT_CLOSE
    Z1_B3 --> EV_MIN
    MAIN --> EV_MIN
    EXT_API --> EV_MIN --> ACT_MIN
    Z1_B4 --> EV_MAX
    MAIN --> EV_MAX --> ACT_MAX

    %% 主题/样式间接依赖 (虚线)
    EXT_THEME -.->|colorText/bg/font| Z1_B1
    EXT_THEME -.->|colorText/bg/font| Z2_CL
    EXT_THEME -.->|colorText/bg/font| Z3_SL
    EXT_THEME -.->|colorText/bg/font| Z4_CB
    EXT_THEME -.->|colorText/bg/font| Z5_BT
    EXT_THEME -.->|colorText/bg/font| Z6_PL
    EXT_THEME -.->|colorText/bg/font| Z7_LB
    EXT_ICON -.->|图标资源| Z1_B1
    EXT_ICON -.->|图标资源| Z1_B2
    EXT_ICON -.->|图标资源| Z1_B3
    EXT_ICON -.->|图标资源| Z1_B4

    %% C++ 后端依赖
    F_ICONBTN -.-> CPP_ICON
    F_LABEL -.-> CPP_THEME
    F_BUTTON -.-> CPP_THEME
    F_CHECK -.-> CPP_THEME
    F_FRAME -.-> CPP_THEME
    F_GROUPBOX -.-> CPP_THEME
    F_RECT -.-> CPP_THEME
    F_ICONBTN -.-> CPP_API

    %% Qt 框架依赖
    C_IMPORT -.-> QQ
    C_IMPORT -.-> QQC
    C_IMPORT -.-> QQL
    C_ROOT --> MAIN

    %% ═══════════════════════════════════════════════
    %% Repeater 数据模型
    %% ═══════════════════════════════════════════════
    subgraph REPEATER["Repeater 静态数据模型"]
        subgraph RM1["参数模型 x8"]
            RM1_D["井底压力/井口压力/回压过高<br/>启用A/B/C阀/启用板A/B"]
        end
        subgraph RM2["按钮模型 x4"]
            RM2_D["参数设置/硬件监控<br/>报告输出/节流阀校准"]
        end
        subgraph RM3["报警模型 x2"]
            RM3_D["气源压力+743KPa<br/>油箱低液位报警"]
        end
    end

    RM1_D --> Z4_RP
    RM2_D --> Z5_RP
    RM3_D --> Z7_RP

    %% ═══════════════════════════════════════════════
    %% 样式 Style
    %% ═══════════════════════════════════════════════
    style CORE fill:#ff6b35,color:#fff,stroke:#c40,stroke-width:3px
    style C_ROOT fill:#ff6b35,color:#fff,stroke:#c40
    style MAIN fill:#4a90d9,color:#fff,stroke:#2a5a8a
    style QT fill:#2d5a1e,color:#af8,stroke:#4a8a2a
    style PROPS fill:#1e5a3d,color:#af8,stroke:#3a8a4a
    style EXT fill:#3a5a1e,color:#dfa,stroke:#5a8a3a
    style CPP fill:#5a3ae0,color:#ddf,stroke:#7a5aff
    style CPP_ICON fill:#8b5cf6,color:#fff,stroke:#6a3ad0
    style CPP_THEME fill:#8b5cf6,color:#fff,stroke:#6a3ad0
    style CPP_API fill:#8b5cf6,color:#fff,stroke:#6a3ad0
    style TREE fill:#1e3a5f,color:#aad,stroke:#4a7ab5
    style Z1 fill:#2e4a2e,color:#ada,stroke:#4a8a4a
    style Z2 fill:#2e3a2e,color:#ada,stroke:#3a6a3a
    style Z3 fill:#5a4a1e,color:#ffd,stroke:#aa8a3a
    style Z4 fill:#3a2e5a,color:#dda,stroke:#7a5aaa
    style Z5 fill:#4a2e3a,color:#fcc,stroke:#aa5a6a
    style Z6 fill:#1e4a4a,color:#add,stroke:#4a9a9a
    style Z7 fill:#4a1e1e,color:#faa,stroke:#aa4a4a
    style BINDINGS fill:#2e2e4a,color:#ccf,stroke:#4a4a8a
    style EVENTS fill:#3a3a5a,color:#ddf,stroke:#5a5a8a
    style ACTIONS fill:#4a2e2e,color:#faa,stroke:#8a4a4a
    style REPEATER fill:#2e1a3a,color:#fda,stroke:#6a3a7a
    style RM1 fill:#3a1a4a,color:#dda,stroke:#6a3aaa
    style RM2 fill:#1a3a4a,color:#add,stroke:#3a7aaa
    style RM3 fill:#4a1a1a,color:#faa,stroke:#9a3a3a
    style LEGEND fill:#1a1a2e,color:#eee,stroke:#333"""

add_diagram(doc, '', unified_mmd)

# ==================== 8. 表格汇总 ====================
doc.add_heading('三、组件依赖清单', level=1)

ct = doc.add_table(rows=14, cols=4, style='Light Grid Accent 1')
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['序号', '组件', '来源', '用途']):
    ct.rows[0].cells[i].text = h
    for pp in ct.rows[0].cells[i].paragraphs:
        pp.runs[0].font.bold = True
for i, d in enumerate([
    ['1', 'MosRectangle', 'MosuiBasic', '外边框 + 参数状态指示灯'],
    ['2', 'MosSpace', 'MosuiBasic', '布局容器 RowLayout/ColumnLayout'],
    ['3', 'MosFrame', 'MosuiBasic', '分区框架: 按钮/状态/参数/控制/报警'],
    ['4', 'MosIconButton', 'MosuiBasic', '启动/退出/最小化/最大化'],
    ['5', 'MosLabel', 'MosuiBasic', '时钟/状态/压力/温度/加热/报警'],
    ['6', 'MosCheckBox', 'MosuiBasic', '8项参数只读状态指示'],
    ['7', 'MosButton', 'MosuiBasic', '参数设置/硬件监控/报告输出/节流阀校准'],
    ['8', 'MosGroupBox', 'MosuiBasic', '液压站数据分组框'],
    ['9', 'MosIcon', 'MosuiBasic', 'UserOutlined/PowerswitchOutlined 等图标'],
    ['10', 'MosTheme', 'MosuiBasic', 'colorSplit/Primary.radiusPrimary 主题'],
    ['11', 'MosRadius', 'MosuiBasic', '{all, topLeft, topRight...} 圆角配置'],
    ['12', 'MosApi', 'MosuiBasic', 'setWindowState 窗口API'],
]):
    for j, v in enumerate(d):
        ct.rows[i+1].cells[j].text = v

doc.add_heading('四、属性与外部依赖', level=1)
pt = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
pt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['依赖项', '类型', '使用方式', '说明']):
    pt.rows[0].cells[i].text = h
    for pp in pt.rows[0].cells[i].paragraphs:
        pp.runs[0].font.bold = True
for i, d in enumerate([
    ['yali_pressure', 'real', 'MosLabel.text', '压力 MPa, toFixed(1)'],
    ['wendu_temperature', 'real', 'MosLabel.text', '温度 C, toFixed(1)'],
    ['heatStatus', 'int(0/1/2)', 'Rectangle.color + MosLabel.text', '0=停止灰 1=运行绿 2=故障红'],
    ['plcappwindow', 'Window', 'onClicked', 'close/showNormal/showMaximized'],
    ['MosApi', 'C++ 单例', 'setWindowState', '调用 Windows API 最小化'],
    ['new Date()', 'JS 内置', 'Timer + formatDateTime', '每秒刷新时钟'],
]):
    for j, v in enumerate(d):
        pt.rows[i+1].cells[j].text = v

# ==================== 保存 ====================
output_dir = r'e:\qt\Mosui\Tools'
output_path = os.path.join(output_dir, 'RightDelegateWin_融合依赖图.docx')
doc.save(output_path)
print(f'OK: {output_path}')
