#!/usr/bin/env python3
"""生成 K3Client 发送逻辑详解 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r"E:\qt\Mosui\Tools\K3Client发送逻辑详解.docx"

doc = Document()

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

def add_code(doc, text, font_size=9):
    """添加代码块段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_note(doc, text):
    """添加注释/说明段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    return p

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表后空行
    return table

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('K3Client 发送逻辑详解')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('K3 PLC 通信协议客户端 — 架构、流程与函数手册')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('文件: Mosuiapp/Tools/PLC_APP/PLCbackend/K3Client.cpp / .h\n依赖: MosuiBasic/cpp/tools/MosNetTcp/\n日期: 2026-07-06')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、架构概览',
    '二、协议帧格式',
    '  2.1 帧结构',
    '  2.2 协议常量',
    '  2.3 命令体格式',
    '三、函数详解',
    '  3.1 构造 / 析构',
    '  3.2 单例与 QML 注册',
    '  3.3 属性访问器',
    '  3.4 连接控制 (connectToHost / disconnectFromHost)',
    '  3.5 命令入队 (enqueueCommand / sendNextCommand)',
    '  3.6 帧编码 / 解码 (buildFrame / extractFrame)',
    '  3.7 响应处理 (handleResponse)',
    '  3.8 TCP 回调 (4 个槽函数)',
    '  3.9 公开 API (QML 可调用)',
    '四、完整生命周期时序',
    '五、关键设计决策',
    '六、性能优化点',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 一、架构概览
# ═══════════════════════════════════════════════════════════════
doc.add_heading('一、架构概览', level=1)

doc.add_paragraph(
    'K3Client 是 Mosui 项目中的 PLC 通信客户端，负责与 K3 系列 PLC 通过 TCP 进行数据交互。'
    '它是一个 QML 单例，可在 QML 中直接调用，底层通过 MosNetTcpManager 管理 TCP 连接。'
)

doc.add_heading('1.1 分层架构', level=2)

add_code(doc, '''┌─────────────────────────────────────────────────────────┐
│  QML 层 (用户调用)                                        │
│  dbReadReal / dbReadBit / dbWriteReal / dbWriteBit       │
├─────────────────────────────────────────────────────────┤
│  K3Client (本类 - 协议/队列层)                             │
│  ┌──────────────┐  ┌──────────────┐  ┌───────────────┐  │
│  │ 命令入队       │→│ 逐条发送      │→│ 响应解析/分发  │  │
│  │ enqueueCommand│  │ sendNextCmd  │  │ handleResponse│  │
│  └──────────────┘  └──────┬───────┘  └───────────────┘  │
│                           │ buildFrame (帧编码)           │
├───────────────────────────┼─────────────────────────────┤
│  MosNetTcpManager          │ sendBytes()                  │
│  (TCP 传输层 - 自动重连)    ↓                              │
│  ← dataReceived 信号 → extractFrame (帧解码)              │
└─────────────────────────────────────────────────────────┘''')

doc.add_heading('1.2 核心设计原则', level=2)

make_table(doc,
    ['原则', '说明'],
    [
        ['单例模式', '全局唯一实例，QML 侧直接 K3Client.xxx() 调用，无需手动管理生命周期'],
        ['命令队列', '所有读写操作排队，串行执行，保证 PLC 响应顺序与请求顺序一致'],
        ['帧协议', '$ + hex(命令体) + hex(校验和) + #，校验和为各字节累加和'],
        ['主备双地址', '主地址连不上时自动切换到备地址，提高可用性'],
        ['断线保留队列', 'TCP 断开时不丢弃积压命令，重连后自动恢复发送'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('1.3 类继承与依赖', level=2)

add_code(doc, '''K3Client : public QObject
  ├── 依赖 MosNetTcpManager (单例) — TCP 传输
  ├── QML_SINGLETON + QML_NAMED_ELEMENT("K3Client")
  ├── 内部结构体 PendingCommand — 队列元素
  └── 使用 std::deque<PendingCommand> — 命令队列''')

# ═══════════════════════════════════════════════════════════════
# 二、协议帧格式
# ═══════════════════════════════════════════════════════════════
doc.add_heading('二、协议帧格式', level=1)

doc.add_heading('2.1 帧结构', level=2)

doc.add_paragraph('K3 协议将原始命令体编码为 hex 字符串，加上帧头帧尾和校验和，通过 TCP 发送。')

add_code(doc, '''帧结构:  $  HH  HH  ...  HH  CC  #
         ↑  └─── 命令体 hex ───┘  ↑   ↑
       帧头    每字节→2字符hex    校验 帧尾

示例: 原始命令 [01 04 00 00 01 00]  (6字节)
      → hex  "010400000100"
      → 校验和 = 01+04+00+00+01+00 = 06
      → 帧   "$01040000010006#"''')

doc.add_heading('2.2 协议常量', level=2)

make_table(doc,
    ['常量名', '值', '含义'],
    [
        ['kDeviceId', '1', '本机设备地址'],
        ['kCmdReadReal', '4', '读取浮点寄存器'],
        ['kCmdWriteReal', '5', '写入浮点寄存器'],
        ['kCmdReadBit', '6', '读取位寄存器'],
        ['kCmdWriteBit', '7', '写入位寄存器'],
        ['kFrameStart', "'$'", '帧头标识'],
        ['kFrameEnd', "'#'", '帧尾标识'],
        ['kMaxBuffer', '4096', '接收缓冲区最大字节数'],
    ],
    col_widths=[4, 3, 9]
)

doc.add_heading('2.3 命令体格式', level=2)

doc.add_paragraph('所有多字节整数均使用小端序 (Little Endian)。', style='List Bullet')

doc.add_heading('2.3.1 ReadReal / ReadBit 命令 (6 字节)', level=3)

add_code(doc, '''偏移  0       1       2    3       4    5
     [DevID] [CmdID] [AddrLo] [AddrHi] [SizeLo] [SizeHi]

     Addr = dbNumber + start - 1''')

doc.add_heading('2.3.2 WriteReal 命令 (9 字节)', level=3)

add_code(doc, '''偏移  0       1       2    3       4    5    6  7  8  9
     [DevID] [CmdID] [AddrLo] [AddrHi] [1] [0] [Float × 4 bytes]
                                          └─Count─┘''')

doc.add_heading('2.3.3 WriteBit 命令 (6 字节)', level=3)

add_code(doc, '''偏移  0       1       2    3       4        5
     [DevID] [CmdID] [AddrLo] [AddrHi] [BitPos] [Value: 0/1]''')

doc.add_heading('2.3.4 响应体格式', level=3)

make_table(doc,
    ['命令', '响应格式', '备注'],
    [
        ['ReadReal', '[Dev(1)][Cmd(1)][ByteCount(2)][Float×N]', 'ByteCount/4 = 浮点个数'],
        ['ReadBit', '[Dev(1)][Cmd(1)][ByteCount(2)][Byte×N]', 'ByteCount = 数据字节数'],
        ['WriteReal', '[Dev(1)][Cmd(1)]', '仅返回确认'],
        ['WriteBit', '[Dev(1)][Cmd(1)]', '仅返回确认'],
    ],
    col_widths=[3, 8, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 三、函数详解
# ═══════════════════════════════════════════════════════════════
doc.add_heading('三、函数详解', level=1)

# 3.1
doc.add_heading('3.1 构造 / 析构', level=2)

doc.add_heading('K3Client::K3Client(QObject *parent)', level=3)
doc.add_paragraph('构造函数，初始化 TCP 连接参数并绑定信号槽。')

add_code(doc, '''K3Client::K3Client(QObject *parent) : QObject(parent)
{
    m_tcp = MosNetTcpManager::instance();

    m_tcp->setHost(m_host);              // 默认 192.168.100.18
    m_tcp->setPort(m_port);              // 默认 50002
    m_tcp->setAutoReconnect(true);       // ★ 启用自动重连
    m_tcp->setReconnectInterval(3000);   // 3 秒重试间隔
    m_tcp->setTcpNoDelay(true);          // 禁用 Nagle，低延迟
    m_tcp->setLazyDecode(true);          // ★ 延迟解码，仅用 raw bytes

    // 绑定 4 个信号
    connect(m_tcp, &MosNetTcpManager::connected,     → onTcpConnected)
    connect(m_tcp, &MosNetTcpManager::disconnected,  → onTcpDisconnected)
    connect(m_tcp, &MosNetTcpManager::dataReceived,  → onTcpDataReceived)
    connect(m_tcp, &MosNetTcpManager::errorOccurred, → onTcpError)
}''')

add_note(doc, '注: lazyDecode = true 使 TCP 层不做文本/hex 解码，直接传递原始字节，减少不必要的 CPU 开销。')

doc.add_heading('K3Client::~K3Client()', level=3)
doc.add_paragraph('默认析构函数 (= default)。由于 m_tcp 是外部单例，不负责其生命周期。')

# 3.2
doc.add_heading('3.2 单例与 QML 注册', level=2)

doc.add_heading('static K3Client *instance()', level=3)
doc.add_paragraph('返回全局唯一实例（函数内 static 局部变量，C++11 线程安全）。')

add_code(doc, '''static K3Client *instance()
{
    static K3Client ins;
    return &ins;
}''')

doc.add_heading('static K3Client *create(QQmlEngine *, QJSEngine *)', level=3)
doc.add_paragraph('QML 单例工厂方法。QML 引擎在首次访问 K3Client 时自动调用，设置 CppOwnership 防止 GC 回收。')

add_code(doc, '''static K3Client *create(QQmlEngine *, QJSEngine *)
{
    auto *c = instance();
    QQmlEngine::setObjectOwnership(c, QQmlEngine::CppOwnership);
    return c;
}''')

# 3.3
doc.add_heading('3.3 属性访问器', level=2)

doc.add_paragraph('K3Client 暴露 5 个 Q_PROPERTY，对应 5 个 getter 和 4 个 setter（isConnected 只读）。')

make_table(doc,
    ['属性', '类型', '默认值', '读写', '信号'],
    [
        ['isConnected', 'bool', 'false', '只读', 'isConnectedChanged'],
        ['host', 'QString', '192.168.100.18', '读写', 'hostChanged'],
        ['port', 'int', '50002', '读写', 'portChanged'],
        ['secondaryHost', 'QString', '192.168.100.18', '读写', 'secondaryHostChanged'],
        ['secondaryPort', 'int', '50000', '读写', 'secondaryPortChanged'],
    ],
    col_widths=[3.5, 2.5, 4, 2, 4]
)

doc.add_paragraph('setter 内部逻辑统一：')
add_code(doc, '''void K3Client::setHost(const QString &h)
{
    if (m_host == h) return;    // 值未变 → 跳过
    m_host = h;
    m_tcp->setHost(h);          // 同步更新 TCP 层
    emit hostChanged();         // 通知 QML
}''')

# 3.4
doc.add_heading('3.4 连接控制', level=2)

doc.add_heading('connectToHost()', level=3)
doc.add_paragraph('QML 手动触发连接。会先断开旧连接，清空队列和缓冲区，然后尝试主地址。')

add_code(doc, '''void connectToHost()
{
    1. if (m_tcp->isConnected())  → m_tcp->disconnectFromHost()
    2. m_queue.clear()              // 清空命令队列
    3. m_commandInFlight = false
    4. m_readBuffer.clear()
    5. tryConnectPrimary()
}''')

doc.add_heading('tryConnectPrimary()', level=3)
add_code(doc, '''void tryConnectPrimary()
{
    m_tcp->setHost(m_host);       // 设为主地址 (默认 192.168.100.18:50002)
    m_tcp->setPort(m_port);
    m_tcp->connectToHost();       // 发起 TCP 连接(异步)
}''')

doc.add_heading('tryConnectSecondary()', level=3)
add_code(doc, '''void tryConnectSecondary()
{
    m_tcp->setHost(m_secondaryHost);  // 设为备地址 (默认 192.168.100.18:50000)
    m_tcp->setPort(m_secondaryPort);
    m_tcp->connectToHost();
}''')

add_note(doc, '★ 主备切换仅在 onTcpError 中触发。主地址连接失败时自动尝试备地址，无需用户干预。')

doc.add_heading('disconnectFromHost()', level=3)
add_code(doc, '''void disconnectFromHost()
{
    m_queue.clear();
    m_commandInFlight = false;
    m_tcp->disconnectFromHost();
}''')

# 3.5
doc.add_heading('3.5 命令入队 (核心流程入口)', level=2)

doc.add_heading('enqueueCommand(quint8 cmdId, QByteArray &&cmd, int dbNumber, int start, bool isWrite)', level=3)
doc.add_paragraph('统一的命令入队入口，消除 4 个公开 API 的重复代码。这是整个发送流程的起点。')

add_code(doc, '''void enqueueCommand(cmdId, cmd, dbNumber, start, isWrite)
{
    1. wasIdle = (m_queue.empty() && !m_commandInFlight)
       // 记录入队前是否空闲

    2. 构造 PendingCommand:
       PendingCommand {
           rawCommand = std::move(cmd)  // ★ 移动语义，零拷贝
           cmdId      = kCmdReadReal / kCmdReadBit / ...
           dbNumber   = 数据库块号
           start      = 起始偏移
           isWrite    = 是否写操作
       }

    3. m_queue.push_back(std::move(pc))  // 追加到队尾 (deque)

    4. if (wasIdle && m_connected)  → sendNextCommand()
       // 空闲时立即发送；繁忙时排队等待
}''')

add_note(doc, '关键设计: 空闲时立即发送，繁忙时排队。保证 PLC 收到的命令顺序与调用顺序严格一致。')

doc.add_heading('sendNextCommand()', level=3)
doc.add_paragraph('从队列头部取出下一条命令并发送。仅发送，不移除队列（移除在收到响应后）。')

add_code(doc, '''void sendNextCommand()
{
    1. if (m_queue.empty() || !m_connected)
       {
           m_commandInFlight = false;
           return;  // 无命令或未连接 → 什么都不做
       }

    2. m_commandInFlight = true  // ★ 标记正在发送，防止并发

    3. const PendingCommand &pc = m_queue.front()
       // 取队首，不移除

    4. QByteArray frame = buildFrame(pc.rawCommand)
       // 编码为 "$hexhex...#"

    5. m_tcp->sendBytes(frame)
       // 交给 TCP 层发送
}''')

add_note(doc, '★ 为什么发送后不移除队列？因为命令尚未被确认。如果在发送后立即断线，命令仍在队列中，重连后 onTcpConnected → sendNextCommand 会自动重发。')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('PendingCommand 结构体定义:')
run.font.bold = True

add_code(doc, '''struct PendingCommand {
    QByteArray rawCommand;   // 原始命令字节
    quint8      cmdId;       // 命令类型 (4/5/6/7)
    int         dbNumber;    // DB 块号
    int         start;       // 起始偏移
    bool        isWrite;     // 是否写操作
};''')

# 3.6
doc.add_heading('3.6 帧编码 / 解码', level=2)

doc.add_heading('buildFrame(const QByteArray &command)', level=3)

doc.add_paragraph('将原始命令字节编码为 K3 协议帧。使用编译时查表 + 直接指针写入，零临时分配。')

add_code(doc, '''QByteArray buildFrame(command)
{
    frameSize = 1(帧头$) + cmdSize×2(hex编码) + 2(校验hex) + 1(帧尾#)

    1. frame.resize(frameSize)  // 预分配精确大小
    2. char *p = frame.data()   // 直接指针操作

    3. *p++ = '$'               // 写帧头

    4. // ★ 核心循环: 查表编码 + 累加校验和，一次完成
       cks = 0
       for (i = 0; i < cmdSize; i++):
           b = command[i]
           cks += b
           hex = s_hex[b]       // 查 256×3 编译时常量表
           *p++ = hex[0]
           *p++ = hex[1]

    5. cksHex = s_hex[cks]      // 校验和也查表
       *p++ = cksHex[0]
       *p++ = cksHex[1]

    6. *p = '#'                 // 写帧尾

    return frame
}''')

add_note(doc, '★ 性能优化: s_hex[256][3] 是 768 字节编译时常量数组，避免了每字节调用 QByteArray::number().rightJustified().toUpper() 产生的大量临时对象分配。')

doc.add_heading('extractFrame(QByteArray &buffer, QByteArray &frameData)', level=3)

doc.add_paragraph('从接收缓冲区中提取一个完整帧，解码为原始字节，并验证校验和。')

add_code(doc, '''bool extractFrame(buffer, out frameData)
{
    1. 查找 '$' 位置 startIdx:
       - 未找到 → 清空 buffer，返回 false
       - 不在开头 → buffer.remove(0, startIdx)

    2. 查找 '#' 位置 endIdx:
       - 未找到 → 保留 buffer 等更多数据，返回 false
       - hexLen = endIdx - 1    (排除 $ 和 #)

    3. 合法性校验:
       - hexLen < 4 或奇数 → 丢弃，返回 false

    4. 解码:
       binary = QByteArray::fromHex(
           QByteArray::fromRawData(buffer.constData() + 1, hexLen)
       )
       // ★ fromRawData 零拷贝引用，fromHex 一次分配

    5. buffer.remove(0, endIdx + 1)  // 移除已处理的帧

    6. 校验和验证:
       dataLen = binary.size() - 1
       cks = 0
       for (i = 0; i < dataLen; i++):
           cks += binary[i]
       if (cks != binary[dataLen]) → 丢弃，返回 false

    7. frameData = binary.left(dataLen)  // 返回不含校验和的数据
       // left() 隐式共享，不实际拷贝
       return true
}''')

doc.add_paragraph()
make_table(doc,
    ['步骤', '失败处理', '目的'],
    [
        ['未找到 $', '清空 buffer', '丢弃无效数据，等待新帧起始'],
        ['未找到 #', '保留 buffer', '帧不完整，等待后续 TCP 数据到达'],
        ['hexLen 非法', '跳过该帧', '防止畸形数据导致崩溃'],
        ['校验和不匹配', '丢弃该帧', '数据损坏，不向上传递'],
    ],
    col_widths=[4, 3.5, 8.5]
)

# 3.7
doc.add_heading('3.7 响应处理', level=2)

doc.add_heading('handleResponse(QByteArray &&response)', level=3)

doc.add_paragraph(
    '处理一条完整协议响应。负责校验响应合法性，解析数据，发射相应信号，并驱动队列继续。'
    '这是连接"接收"和"发送"的核心枢纽。'
)

add_code(doc, '''void handleResponse(response)
{
    ── 前提条件 ──
    0. if (!m_commandInFlight || m_queue.empty()) → return
       // 没有待处理的命令，忽略
    1. pc = std::move(m_queue.front())  // ★ 移动语义，不拷贝 QByteArray
       m_queue.pop_front()              // 移出队列
       m_commandInFlight = false

    ── 基本校验 ──
    2. if (response.size() < 2)  → emit errorOccurred; sendNextCmd
    3. rspDev = response[0];  rspCmd = response[1]
       if (rspDev != kDeviceId || rspCmd != pc.cmdId)
           → emit errorOccurred; sendNextCmd

    ── 写入确认 ──
    4. if (pc.isWrite)
       {
           emit writeCompleted(pc.dbNumber, pc.start, pc.cmdId);
           sendNextCommand();
           return;
       }

    ── 读取解析 ──
    5. switch (pc.cmdId)
       {
       case kCmdReadReal:
           byteCount = qFromLittleEndian<quint16>(raw + 2)
           count = byteCount / 4
           values = QVector<float>(count)    // ★ 单次分配
           for i in count:
               values[i] = qFromLittleEndian<float>(raw + 4 + i*4)
           emit realDataReceived(dbNumber, start, values)

       case kCmdReadBit:
           byteCount = qFromLittleEndian<quint16>(raw + 2)
           values = QVector<quint8>(byteCount)
           for i in byteCount:
               values[i] = response[4 + i]
           emit bitDataReceived(dbNumber, start, values)
       }

    6. sendNextCommand()  // ★ 继续处理队列中的下一条
}''')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('信号发送对照表:')
run.font.bold = True

make_table(doc,
    ['命令类型', '发射的信号', '数据载体'],
    [
        ['ReadReal', 'realDataReceived(dbNumber, start, values)', 'QVector<float>'],
        ['ReadBit', 'bitDataReceived(dbNumber, start, rawBytes)', 'QVector<quint8>'],
        ['WriteReal/WriteBit', 'writeCompleted(dbNumber, start, cmdId)', '无数据'],
        ['异常', 'errorOccurred(message)', 'QString'],
    ],
    col_widths=[3.5, 8, 4.5]
)

add_note(doc, '★ 无论解析成功还是失败，handleResponse 最终都会调用 sendNextCommand() 驱动队列前进。这保证了即使某条命令响应异常，队列也不会永久阻塞。')

# 3.8
doc.add_heading('3.8 TCP 回调 (4 个槽函数)', level=2)

doc.add_heading('onTcpConnected()', level=3)
add_code(doc, '''void onTcpConnected()
{
    1. if (!m_connected):
       m_connected = true → emit isConnectedChanged()
    2. m_readBuffer.clear()        // 丢弃旧连接残留
    3. if (!m_queue.empty() && !m_commandInFlight):
       sendNextCommand()           // ★ 重连后自动恢复积压命令
}''')

doc.add_heading('onTcpDisconnected()', level=3)
add_code(doc, '''void onTcpDisconnected()
{
    1. m_commandInFlight = false   // 当前命令作废
    2. m_readBuffer.clear()        // 丢弃不完整帧
    3. if (m_connected):
       m_connected = false → emit isConnectedChanged()

    // ★ 不清空 m_queue！断线保留队列，重连后自动继续
}''')

add_note(doc, '★ 关键设计: 断线时保留队列而非清空，是可靠性的核心。如果清空队列，QML 层不知道哪些请求被丢弃，需要重新发起。保留队列让断线对上层透明。')

doc.add_heading('onTcpDataReceived(const QByteArray &data, ...)', level=3)
add_code(doc, '''void onTcpDataReceived(data, _, _)
{
    1. m_readBuffer.append(data)   // 追加新数据

    2. while (extractFrame(m_readBuffer, frameBody)):
       // 循环提取所有完整帧
       if (!frameBody.isEmpty()):
           handleResponse(frameBody)  // 处理每一帧

    3. if (m_readBuffer.size() > 4096):
       m_readBuffer = m_readBuffer.right(2048)
       // 缓冲区过大 → 截断保留后半，防止内存无限增长
}''')

doc.add_heading('onTcpError(const QString &msg)', level=3)
add_code(doc, '''void onTcpError(msg)
{
    1. if (!m_connected && m_tcp->host() == m_host):
       // 主地址连不上 → ★ 自动切换备地址
       tryConnectSecondary()
       return   // 不向外 emit 错误

    2. emit errorOccurred(msg)  // 其他错误才通知上层
}''')

# 3.9
doc.add_heading('3.9 公开 API (QML 可调用)', level=2)

doc.add_paragraph('4 个 Q_INVOKABLE 方法，均可从 QML 直接调用。内部统一委托给 enqueueCommand。')

doc.add_heading('dbReadReal(int dbNumber, int start, int size)', level=3)
doc.add_paragraph('读取 DB 块中的浮点寄存器。')
add_code(doc, '''QML: K3Client.dbReadReal(dbNumber, start, size)
命令: [01] [04] [AddrLo] [AddrHi] [SizeLo] [SizeHi]
      Addr = dbNumber + start - 1
结果: 收到 realDataReceived(dbNumber, start, QVector<float>) 信号''')

doc.add_heading('dbReadBit(int dbNumber, int start, int size)', level=3)
doc.add_paragraph('读取 DB 块中的位寄存器。')
add_code(doc, '''QML: K3Client.dbReadBit(dbNumber, start, size)
命令: [01] [06] [AddrLo] [AddrHi] [SizeLo] [SizeHi]
结果: 收到 bitDataReceived(dbNumber, start, QVector<quint8>) 信号''')

doc.add_heading('dbWriteReal(float value, int dbNumber, int start)', level=3)
doc.add_paragraph('写入单个浮点到 DB 寄存器。')
add_code(doc, '''QML: K3Client.dbWriteReal(1.5, dbNumber, start)
命令: [01] [05] [AddrLo] [AddrHi] [01] [00] [Float×4 LE]
结果: 收到 writeCompleted(dbNumber, start, kCmdWriteReal) 信号''')

doc.add_heading('dbWriteBit(bool value, int dbNumber, int start, int bit)', level=3)
doc.add_paragraph('写入单个位到 DB 寄存器。')
add_code(doc, '''QML: K3Client.dbWriteBit(true, dbNumber, start, bit)
命令: [01] [07] [AddrLo] [AddrHi] [bit] [0/1]
结果: 收到 writeCompleted(dbNumber, start, kCmdWriteBit) 信号''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 四、完整生命周期时序
# ═══════════════════════════════════════════════════════════════
doc.add_heading('四、完整生命周期时序', level=1)

doc.add_heading('4.1 正常读取流程', level=2)

add_code(doc, '''用户 QML:  K3Client.dbReadReal(1, 0, 10)
         │
         ▼
    enqueueCommand(kCmdReadReal, cmd, 1, 0, false)
         │
         ├─ 队列空 + 已连接 → sendNextCommand()
         │                          │
         │                          ▼
         │               buildFrame([01 04 00 00 0A 00])
         │                          │
         │                          ▼
         │               m_tcp->sendBytes("$010400000A000F#")
         │
    ── 等待 PLC 响应 ──
         │
         ▼
    onTcpDataReceived(raw_bytes)
         │
         ▼
    extractFrame → 解码出 [01 04 14 00 00 80 3F ...]
         │
         ▼
    handleResponse(response)
         │
         ├─ 校验 DevID=1, CmdID=4 ✓
         ├─ 解析 ByteCount=20 → 5 个 float
         ├─ emit realDataReceived(1, 0, [1.0, 2.0, 3.0, 4.0, 5.0])
         │
         ▼
    sendNextCommand()  → 继续处理队列中的下一条''')

doc.add_heading('4.2 连续多条命令流程', level=2)

add_code(doc, '''QML 连续调用:
  dbReadReal(1, 0, 10)     ← 第 1 条
  dbWriteReal(5.0, 2, 0)   ← 第 2 条
  dbReadBit(3, 0, 8)       ← 第 3 条

时间线:
  T0: enqueueCommand(#1) → 空闲 → sendNextCommand → 发送帧#1
  T1: enqueueCommand(#2) → 繁忙 → 仅入队
  T2: enqueueCommand(#3) → 繁忙 → 仅入队
  T3: handleResponse(#1) → emit realDataReceived → sendNextCommand → 发送帧#2
  T4: handleResponse(#2) → emit writeCompleted → sendNextCommand → 发送帧#3
  T5: handleResponse(#3) → emit bitDataReceived → sendNextCommand → 队列空，停止

队列状态变化:
  [#1] → [#1, #2] → [#1, #2, #3] → [#2, #3] → [#3] → []''')

doc.add_heading('4.3 断线重连流程', level=2)

add_code(doc, '''场景: 发送中 TCP 断线

  队列: [cmd#1, cmd#2, cmd#3]   ← #1 正在发送中
         │
  TCP 断开
         │
         ▼
  onTcpDisconnected()
    → m_commandInFlight = false
    → m_connected = false
    → emit isConnectedChanged()
    → ★ m_queue 保留 [cmd#1, cmd#2, cmd#3]
         │
  (3 秒后)...MosNetTcpManager 自动重连...
         │
         ▼
  onTcpConnected()
    → m_connected = true
    → emit isConnectedChanged()
    → m_queue 非空 && !m_commandInFlight
    → sendNextCommand()  ← ★ 自动重发 cmd#1

  正常完成 cmd#1, cmd#2, cmd#3...''')

doc.add_heading('4.4 主备切换流程', level=2)

add_code(doc, '''场景: 主地址 (50002) 连不上

  tryConnectPrimary() → TCP 连接 192.168.100.18:50002
         │
         ▼ (连接失败)
  onTcpError(msg)
    → !m_connected ✓
    → m_tcp->host() == m_host ✓
    → tryConnectSecondary()  ← ★ 自动切备地址
         │
         ▼
  TCP 连接 192.168.100.18:50000
         │
         ├─ 成功 → onTcpConnected → 正常通信
         └─ 失败 → onTcpError → emit errorOccurred(msg)
                   // 备地址也失败，向外通知错误''')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 五、关键设计决策
# ═══════════════════════════════════════════════════════════════
doc.add_heading('五、关键设计决策', level=1)

make_table(doc,
    ['决策', '理由', '影响'],
    [
        ['命令串行队列', 'PLC 通常不支持并发请求；保证响应顺序与请求严格一致', '简单可靠，但高频率读写可能积压'],
        ['断线保留队列', '避免数据丢失；TCP 断开对上层 QML 透明', '重连后自动恢复，无需 QML 重发'],
        ['主备双地址', '主 PLC 故障或网络问题时自动切换到备机', '提高系统可用性，无需人工干预'],
        ['发送后不移除队列', '防止"已发送未确认"时断线导致命令丢失', '命令只有在收到有效响应后才出队'],
        ['std::deque 而非 QList', 'deque 对头删 (pop_front) 更高效', '队列操作零额外开销'],
        ['移动语义传递 QByteArray', '避免命令体在入队/出队时的数据拷贝', '大命令体效率提升显著'],
        ['QVector 替代 QVariantList', 'QML 可直接当作 JS 数组遍历；无逐元素 QVariant 堆分配', '内存效率高，约减少 80% 响应处理开销'],
        ['单例模式', '全局唯一 PLC 连接；避免多实例导致的资源竞争', 'QML 中统一入口 K3Client.xxx()'],
    ],
    col_widths=[4, 7, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 六、性能优化点
# ═══════════════════════════════════════════════════════════════
doc.add_heading('六、性能优化点', level=1)

doc.add_paragraph('代码中标有 ★ 的关键优化：')

make_table(doc,
    ['#', '优化项', '位置', '技术手段', '效果'],
    [
        ['1', 'Hex 编码查表', 'buildFrame',
         's_hex[256][3] 编译时常量数组，256×3=768 字节',
         '消除每字节 QByteArray::number().rightJustified().toUpper() 的临时分配'],
        ['2', '直接指针写入', 'buildFrame',
         'char *p = frame.data() 逐字节写入',
         '一次 resize 后零分配；编码+校验和在一轮循环完成'],
        ['3', 'fromRawData 零拷贝', 'extractFrame',
         'QByteArray::fromRawData 引用原始 buffer 内存',
         '避免 mid() 产生的临时 QByteArray 拷贝'],
        ['4', '统一入队入口', 'enqueueCommand',
         '4 个公开 API 委托给同一方法',
         '消除 dbReadReal/dbReadBit/dbWriteReal/dbWriteBit 的重复代码'],
        ['5', '移动语义', 'enqueueCommand / handleResponse',
         'std::move(cmd) / std::move(pc)',
         '命令体在队列间传递无数据拷贝'],
        ['6', 'QVector 信号', 'handleResponse',
         'QVector<float> / QVector<quint8> 替代 QVariantList',
         '单次分配，QML 直接当 JS 数组用'],
        ['7', '延迟解码', '构造',
         'MosNetTcpManager::setLazyDecode(true)',
         'TCP 层不做文本/hex 转换，直接传递 raw bytes'],
    ],
    col_widths=[1, 3.5, 3, 5.5, 5]
)

# ── 结尾 ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── 保存 ──
doc.save(OUTPUT)
print(f"Document saved: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
